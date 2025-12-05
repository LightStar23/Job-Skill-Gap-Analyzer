# app.py
from flask import Flask, render_template, request, jsonify, redirect, url_for, flash
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from models import db, bcrypt, User, Analysis, CoachingPlan, MilestoneProgress
from skill_analyzer import SkillGapAnalyzer
from career_coach import AICareerCoach
import json
import os
from datetime import datetime
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader
from docx import Document

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-secret-key-here-change-in-production'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///job_analyzer.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(__file__), 'uploads')

# Create uploads folder if it doesn't exist
if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

# Allowed file extensions
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx', 'txt'}

# Initialize extensions
db.init_app(app)
bcrypt.init_app(app)

# Login manager setup
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'
login_manager.login_message = 'Please log in to access this page.'

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# ===== FILE HANDLING UTILITIES =====

def allowed_file(filename):
    """Check if file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_file(filepath):
    """Extract text from PDF, DOCX, DOC, or TXT files"""
    try:
        filename = secure_filename(os.path.basename(filepath))
        file_ext = filename.rsplit('.', 1)[1].lower()
        
        text = ""
        
        if file_ext == 'pdf':
            # Extract text from PDF
            try:
                with open(filepath, 'rb') as file:
                    pdf_reader = PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
            except Exception as e:
                return None, f"Error reading PDF: {str(e)}"
        
        elif file_ext in ['docx']:
            # Extract text from DOCX
            try:
                doc = Document(filepath)
                for para in doc.paragraphs:
                    text += para.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " "
                    text += "\n"
            except Exception as e:
                return None, f"Error reading DOCX: {str(e)}"
        
        elif file_ext in ['doc']:
            # For .doc files, we can try to read as text or use python-docx
            # For now, we'll ask user to convert to DOCX or use TXT
            return None, "Please save your .doc file as .docx format for better compatibility"
        
        elif file_ext == 'txt':
            # Read plain text file
            try:
                with open(filepath, 'r', encoding='utf-8') as file:
                    text = file.read()
            except UnicodeDecodeError:
                try:
                    with open(filepath, 'r', encoding='latin-1') as file:
                        text = file.read()
                except Exception as e:
                    return None, f"Error reading TXT file: {str(e)}"
        
        if not text.strip():
            return None, "The file appears to be empty or could not be read"
        
        return text, None
    
    except Exception as e:
        return None, f"Error processing file: {str(e)}"

# Initialize our analyzer
analyzer = SkillGapAnalyzer()
coach = AICareerCoach()

# Routes
@app.route('/')
def index():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    return render_template('index.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    
    if request.method == 'POST':
        email = request.form.get('email')
        first_name = request.form.get('first_name')
        last_name = request.form.get('last_name')
        password = request.form.get('password')
        confirm_password = request.form.get('confirm_password')
        
        # Basic validation
        if password != confirm_password:
            flash('Passwords do not match!', 'error')
            return render_template('register.html')
        
        if User.query.filter_by(email=email).first():
            flash('Email already registered!', 'error')
            return render_template('register.html')
        
        # Create new user
        user = User(
            email=email,
            first_name=first_name,
            last_name=last_name
        )
        user.set_password(password)
        
        db.session.add(user)
        db.session.commit()
        
        flash('Registration successful! Please log in.', 'success')
        return redirect(url_for('login'))
    
    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')
        remember = bool(request.form.get('remember'))
        
        user = User.query.filter_by(email=email).first()
        
        if user and user.check_password(password):
            login_user(user, remember=remember)
            next_page = request.args.get('next')
            flash(f'Welcome back, {user.first_name}!', 'success')
            return redirect(next_page or url_for('dashboard'))
        else:
            flash('Invalid email or password!', 'error')
    
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash('You have been logged out.', 'info')
    return redirect(url_for('index'))

@app.route('/dashboard')
@login_required
def dashboard():
    # Get user's recent analyses
    recent_analyses = Analysis.query.filter_by(user_id=current_user.id)\
        .order_by(Analysis.created_at.desc())\
        .limit(5)\
        .all()
    return render_template('dashboard.html', analyses=recent_analyses)

@app.route('/upload-file', methods=['POST'])
@login_required
def upload_file():
    """Handle file upload for resume or job description"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        file_type = request.form.get('file_type', '')  # 'resume' or 'job_description'
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file type. Allowed: PDF, DOCX, DOC, TXT'}), 400
        
        # Save file temporarily
        filename = secure_filename(file.filename)
        timestamp = int(datetime.utcnow().timestamp())
        filename = f"{timestamp}_{filename}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        file.save(filepath)
        
        # Extract text from file
        text, error = extract_text_from_file(filepath)
        
        if error:
            os.remove(filepath)
            return jsonify({'error': error}), 400
        
        # Clean up the temporary file
        try:
            os.remove(filepath)
        except:
            pass
        
        return jsonify({
            'success': True,
            'text': text,
            'file_type': file_type,
            'filename': file.filename
        }), 200
    
    except Exception as e:
        return jsonify({'error': f'File upload failed: {str(e)}'}), 500

@app.route('/analyze', methods=['POST'])
@login_required
def analyze():
    try:
        job_description = request.form.get('job_description', '')
        resume_text = request.form.get('resume_text', '')
        job_title = request.form.get('job_title', 'Untitled Position')
        company = request.form.get('company', '')
        
        if not job_description or not resume_text:
            return jsonify({'error': 'Please provide both job description and resume text'})
        
        # Perform analysis
        results = analyzer.analyze_gap(job_description, resume_text)
        
        # Save analysis to database
        analysis = Analysis(
            user_id=current_user.id,
            job_title=job_title,
            company=company,
            job_description=job_description,
            resume_text=resume_text,
            match_score=results['match_score'],
            analysis_results=json.dumps(results)
        )
        db.session.add(analysis)
        db.session.commit()
        
        results['analysis_id'] = analysis.id
        
        return jsonify(results)
        
    except Exception as e:
        return jsonify({'error': f'Analysis failed: {str(e)}'})

@app.route('/analysis/<int:analysis_id>')
@login_required
def view_analysis(analysis_id):
    analysis = Analysis.query.filter_by(id=analysis_id, user_id=current_user.id).first_or_404()
    results = json.loads(analysis.analysis_results)
    return render_template('analysis_detail.html', analysis=analysis, results=results)

@app.route('/analyze/new')
@login_required
def analyze_page():
    return render_template('analyze.html')

@app.route('/coach/generate/<int:analysis_id>', methods=['POST'])
@login_required
def generate_coaching_plan_form(analysis_id):
    """Generate coaching plan from analysis and redirect to dashboard."""
    try:
        # Get the analysis
        analysis = Analysis.query.filter_by(id=analysis_id, user_id=current_user.id).first_or_404()
        analysis_results = json.loads(analysis.analysis_results)
        
        # Extract skill gap data
        skill_gap = {
            "score": analysis_results.get("match_score", 0),
            "gaps": [
                {
                    "skill": skill,
                    "importance": 0.8,
                    "current_level": 3,
                    "gap": 5
                }
                for skill in analysis_results.get("missing_skills", [])
            ],
            "strengths": analysis_results.get("matching_skills", [])
        }
        
        # Build user profile
        user_profile = {
            "id": str(current_user.id),
            "name": current_user.first_name or "User",
            "current_role": None,
            "target_role": analysis.job_title,
            "experience_years": 0,
            "availability_per_week_hours": 10,
            "preferred_learning_style": "projects",
            "budget_level": "free",
            "soft_skills": []
        }
        
        # Generate coaching plan using the AI Coach
        coaching_plan = coach.generate_coaching_plan(user_profile, skill_gap)
        
        if "error" in coaching_plan:
            flash(f"Error generating coaching plan: {coaching_plan['error']}", "error")
            return redirect(url_for('analysis_detail', analysis_id=analysis_id))
        
        # Save to database
        db_coaching_plan = CoachingPlan(
            user_id=current_user.id,
            analysis_id=analysis_id,
            target_role=user_profile["target_role"],
            current_role=user_profile["current_role"],
            experience_years=user_profile["experience_years"],
            availability_per_week_hours=user_profile["availability_per_week_hours"],
            preferred_learning_style=user_profile["preferred_learning_style"],
            budget_level=user_profile["budget_level"],
            coaching_data=json.dumps(coaching_plan),
            confidence_score=coaching_plan.get("confidence", 0.5),
            current_milestone_id="m1",
            milestones_completed=json.dumps([])
        )
        
        db.session.add(db_coaching_plan)
        db.session.commit()
        
        flash(f"✨ Coaching plan generated successfully!", "success")
        return redirect(url_for('view_coaching_plan', coaching_plan_id=db_coaching_plan.id))
        
    except Exception as e:
        flash(f"Error generating coaching plan: {str(e)}", "error")
        return redirect(url_for('analysis_detail', analysis_id=analysis_id))

# ===== AI Career Coach Endpoints =====

@app.route('/api/coach/generate', methods=['POST'])
@login_required
def generate_coaching_plan():
    """
    Generate a personalized coaching plan based on skill gap and learner profile.
    
    Expects JSON:
    {
        "analysis_id": <int>,  // Optional: link to existing analysis
        "target_role": "<string>",
        "current_role": "<string|null>",
        "experience_years": <int>,
        "availability_per_week_hours": <int>,
        "preferred_learning_style": "<string>",  // projects, videos, text, mentorship
        "budget_level": "<string>",  // free, low, medium, high
        "location": "<string|null>"
    }
    """
    try:
        data = request.get_json()
        
        # Validate required fields
        if not data.get("target_role"):
            return jsonify({"error": "target_role is required"}), 400
        
        # Get skill gap data from linked analysis if provided
        skill_gap = {"score": 0, "gaps": [], "strengths": []}
        analysis_id = data.get("analysis_id")
        
        if analysis_id:
            analysis = Analysis.query.filter_by(id=analysis_id, user_id=current_user.id).first()
            if analysis:
                analysis_results = json.loads(analysis.analysis_results)
                # Extract gaps and strengths from analysis
                skill_gap = {
                    "score": analysis_results.get("match_score", 0),
                    "gaps": [
                        {
                            "skill": skill,
                            "importance": 0.8,
                            "current_level": 3,
                            "gap": 5
                        }
                        for skill in analysis_results.get("missing_skills", [])
                    ],
                    "strengths": analysis_results.get("matching_skills", [])
                }
        
        # Build user profile
        user_profile = {
            "id": str(current_user.id),
            "name": current_user.first_name or "User",
            "current_role": data.get("current_role"),
            "target_role": data.get("target_role"),
            "experience_years": data.get("experience_years", 0),
            "availability_per_week_hours": data.get("availability_per_week_hours", 10),
            "preferred_learning_style": data.get("preferred_learning_style", "projects"),
            "budget_level": data.get("budget_level", "free"),
            "soft_skills": data.get("soft_skills", [])
        }
        
        # Generate coaching plan
        coaching_plan = coach.generate_coaching_plan(user_profile, skill_gap)
        
        if "error" in coaching_plan:
            return jsonify(coaching_plan), 400
        
        # Save coaching plan to database
        db_coaching_plan = CoachingPlan(
            user_id=current_user.id,
            analysis_id=analysis_id,
            target_role=user_profile["target_role"],
            current_role=user_profile["current_role"],
            experience_years=user_profile["experience_years"],
            availability_per_week_hours=user_profile["availability_per_week_hours"],
            preferred_learning_style=user_profile["preferred_learning_style"],
            budget_level=user_profile["budget_level"],
            coaching_data=json.dumps(coaching_plan),
            confidence_score=coaching_plan.get("confidence", 0.5),
            current_milestone_id="m1",
            milestones_completed=json.dumps([])
        )
        
        db.session.add(db_coaching_plan)
        db.session.commit()
        
        # Return with ID for tracking
        coaching_plan["coaching_plan_id"] = db_coaching_plan.id
        
        return jsonify(coaching_plan), 201
    
    except Exception as e:
        return jsonify({"error": f"Failed to generate coaching plan: {str(e)}"}), 500

@app.route('/api/coach/<int:coaching_plan_id>', methods=['GET'])
@login_required
def get_coaching_plan(coaching_plan_id):
    """Retrieve a saved coaching plan."""
    try:
        plan = CoachingPlan.query.filter_by(id=coaching_plan_id, user_id=current_user.id).first_or_404()
        coaching_data = json.loads(plan.coaching_data)
        completed = json.loads(plan.milestones_completed)
        
        # Calculate progress
        progress = coach.calculate_overall_progress(coaching_data, completed)
        
        return jsonify({
            "id": plan.id,
            "coaching_plan": coaching_data,
            "progress": progress,
            "created_at": plan.created_at.isoformat(),
            "updated_at": plan.updated_at.isoformat()
        }), 200
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/coach/<int:coaching_plan_id>/milestone/<milestone_id>/progress', methods=['POST'])
@login_required
def update_milestone_progress(coaching_plan_id, milestone_id):
    """
    Update progress on a milestone/task.
    
    Expects JSON:
    {
        "task_index": <int>,
        "status": "not_started|in_progress|completed",
        "validation_evidence": {"project_link": "...", "test_score": 85}
    }
    """
    try:
        data = request.get_json()
        plan = CoachingPlan.query.filter_by(id=coaching_plan_id, user_id=current_user.id).first_or_404()
        
        task_index = data.get("task_index", 0)
        status = data.get("status", "in_progress")
        validation = data.get("validation_evidence", {})
        
        # Check if progress record exists
        progress = MilestoneProgress.query.filter_by(
            coaching_plan_id=coaching_plan_id,
            milestone_id=milestone_id,
            task_index=task_index
        ).first()
        
        if not progress:
            progress = MilestoneProgress(
                coaching_plan_id=coaching_plan_id,
                milestone_id=milestone_id,
                task_index=task_index,
                status=status,
                validation_evidence=json.dumps(validation),
                started_at=datetime.utcnow() if status in ["in_progress", "completed"] else None,
                completed_at=datetime.utcnow() if status == "completed" else None
            )
            db.session.add(progress)
        else:
            progress.status = status
            progress.validation_evidence = json.dumps(validation)
            if status == "in_progress" and not progress.started_at:
                progress.started_at = datetime.utcnow()
            if status == "completed":
                progress.completed_at = datetime.utcnow()
        
        db.session.commit()
        
        # Update coaching plan's completed milestones if all tasks in milestone are done
        coaching_data = json.loads(plan.coaching_data)
        completed = json.loads(plan.milestones_completed)
        
        # Check if all tasks in this milestone are completed
        all_tasks_in_milestone = MilestoneProgress.query.filter_by(
            coaching_plan_id=coaching_plan_id,
            milestone_id=milestone_id
        ).all()
        
        all_completed = all(t.status == "completed" for t in all_tasks_in_milestone)
        if all_completed and milestone_id not in completed:
            completed.append(milestone_id)
            plan.milestones_completed = json.dumps(completed)
            plan.updated_at = datetime.utcnow()
            db.session.commit()
        
        return jsonify({
            "status": "success",
            "milestone_id": milestone_id,
            "task_index": task_index,
            "progress_status": status,
            "updated_at": progress.updated_at.isoformat()
        }), 200
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/coach/<int:coaching_plan_id>/download-pdf', methods=['GET'])
@login_required
def download_coaching_plan_pdf(coaching_plan_id):
    """Download coaching plan as PDF."""
    from io import BytesIO
    try:
        # Import reportlab
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
            from reportlab.lib import colors
        except ImportError:
            # If reportlab not available, use simple approach
            return generate_text_pdf(coaching_plan_id)
        
        # Get coaching plan
        plan = CoachingPlan.query.filter_by(id=coaching_plan_id, user_id=current_user.id).first_or_404()
        coaching_data = json.loads(plan.coaching_data)
        completed = json.loads(plan.milestones_completed)
        progress = coach.calculate_overall_progress(coaching_data, completed)
        
        # Create PDF
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        elements = []
        styles = getSampleStyleSheet()
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#667eea'),
            spaceAfter=10,
            alignment=1
        )
        elements.append(Paragraph("🎯 Your Personalized Coaching Plan", title_style))
        elements.append(Spacer(1, 0.3*inch))
        
        # Summary
        elements.append(Paragraph(f"<b>Summary:</b> {coaching_data.get('summary', '')}", styles['Normal']))
        elements.append(Spacer(1, 0.2*inch))
        
        # Progress
        elements.append(Paragraph(f"<b>Overall Progress:</b> {progress.get('progress_percentage', 0)}%", styles['Normal']))
        elements.append(Paragraph(f"<b>Milestones Completed:</b> {len(progress.get('completed_milestones', []))} of {progress.get('total_milestones', 0)}", styles['Normal']))
        elements.append(Spacer(1, 0.2*inch))
        
        # Coach's Message
        elements.append(Paragraph("<b>Coach's Message:</b>", styles['Heading2']))
        elements.append(Paragraph(coaching_data.get('human_message', ''), styles['Normal']))
        elements.append(Spacer(1, 0.3*inch))
        
        # Milestones
        elements.append(Paragraph("Learning Path Milestones", styles['Heading2']))
        for idx, milestone in enumerate(coaching_data.get('prioritized_path', []), 1):
            elements.append(Paragraph(f"<b>Milestone {idx}: {milestone.get('title', '')}</b>", styles['Heading3']))
            elements.append(Paragraph(f"Duration: {milestone.get('time_estimate_weeks')} weeks", styles['Normal']))
            elements.append(Paragraph(f"Why: {milestone.get('why', '')}", styles['Normal']))
            
            skills = ", ".join(milestone.get('skills_targeted', []))
            if skills:
                elements.append(Paragraph(f"Skills: {skills}", styles['Normal']))
            
            elements.append(Spacer(1, 0.1*inch))
        
        # Build PDF
        doc.build(elements)
        buffer.seek(0)
        
        # Return PDF
        return app.response_class(
            response=buffer.getvalue(),
            status=200,
            mimetype="application/pdf",
            headers={"Content-Disposition": "attachment;filename=coaching-plan.pdf"}
        )
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

def generate_text_pdf(coaching_plan_id):
    """Fallback: Generate PDF using simple method when reportlab not available."""
    from io import BytesIO
    
    plan = CoachingPlan.query.filter_by(id=coaching_plan_id, user_id=current_user.id).first_or_404()
    coaching_data = json.loads(plan.coaching_data)
    
    # Create simple text format
    text_content = "COACHING PLAN\n"
    text_content += "=" * 50 + "\n\n"
    text_content += f"Summary: {coaching_data.get('summary', '')}\n\n"
    text_content += f"Coach's Message:\n{coaching_data.get('human_message', '')}\n\n"
    text_content += "MILESTONES:\n"
    
    for idx, milestone in enumerate(coaching_data.get('prioritized_path', []), 1):
        text_content += f"\n{idx}. {milestone.get('title', '')}\n"
        text_content += f"   Duration: {milestone.get('time_estimate_weeks')} weeks\n"
        text_content += f"   Why: {milestone.get('why', '')}\n"
        text_content += f"   Skills: {', '.join(milestone.get('skills_targeted', []))}\n"
    
    # Convert to PDF using browser's print functionality recommendation
    buffer = BytesIO()
    buffer.write(text_content.encode('utf-8'))
    buffer.seek(0)
    
    return app.response_class(
        response=buffer.getvalue(),
        status=200,
        mimetype="text/plain",
        headers={"Content-Disposition": "attachment;filename=coaching-plan.txt"}
    )

@app.route('/api/coach/list', methods=['GET'])
@login_required
def list_coaching_plans():
    """List all coaching plans for the current user."""
    try:
        plans = CoachingPlan.query.filter_by(user_id=current_user.id)\
            .order_by(CoachingPlan.created_at.desc())\
            .all()
        
        result = []
        for plan in plans:
            coaching_data = json.loads(plan.coaching_data)
            completed = json.loads(plan.milestones_completed)
            progress = coach.calculate_overall_progress(coaching_data, completed)
            
            result.append({
                "id": plan.id,
                "target_role": plan.target_role,
                "current_role": plan.current_role,
                "summary": coaching_data.get("summary", ""),
                "confidence": plan.confidence_score,
                "progress_percentage": progress.get("progress_percentage", 0),
                "created_at": plan.created_at.isoformat(),
                "updated_at": plan.updated_at.isoformat()
            })
        
        return jsonify({"plans": result}), 200
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/coach/<int:coaching_plan_id>')
@login_required
def view_coaching_plan(coaching_plan_id):
    """View coaching plan in browser."""
    try:
        plan = CoachingPlan.query.filter_by(id=coaching_plan_id, user_id=current_user.id).first_or_404()
        coaching_data = json.loads(plan.coaching_data)
        completed = json.loads(plan.milestones_completed)
        progress = coach.calculate_overall_progress(coaching_data, completed)
        
        return render_template('coach_plan.html', 
                             coaching_plan=coaching_data,
                             plan_id=plan.id,
                             progress=progress)
    except Exception as e:
        flash(f"Error loading coaching plan: {str(e)}", "error")
        return redirect(url_for('dashboard'))

# Initialize database
with app.app_context():
    db.create_all()

if __name__ == '__main__':
    app.run(debug=True)